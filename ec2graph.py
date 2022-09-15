import json

import boto3
import docx
from docx.shared import Pt, RGBColor, Inches, Mm
from pyasn1.type import tag

from toturial.ec2.ec2link import add_hyperlink

met_data = []
yet_data = []
net_data = []
bet_data = []
pet_data = []
det_data = []


def ec2gr():
    doc = docx.Document()
    ec2client = boto3.client('ec2')
    regions = ec2client.describe_regions(
        AllRegions=False
    )
    for i in regions['Regions']:
        k = (i['RegionName'])
        client = boto3.client('rds', region_name=k)
        response = client.describe_db_instances()
        for rrr in response['DBInstances']:
            run = rrr['DBInstanceIdentifier']
            rrr['region'] = k
            met = ["AWS/RDS", "DatabaseConnections", "DBInstanceIdentifier", run,
                   {"period": 300, "stat": "SampleCount"}]
            # met_data.append(met)
            z = {
                "metrics": [met],
                "legend": {
                    "position": "bottom"
                },
                "region": k,
                "liveData": False,
                "title": "DatabaseConnections: SampleCount",
                "view": "timeSeries",
                "stacked": False
            }
            v = json.dumps(z)
            client = boto3.client('cloudwatch', region_name=k)
            response = client.get_metric_widget_image(MetricWidget=v)

            img = response['MetricWidgetImage']
            text_binary = img
            image = open('DatabaseConnections' + k + '.png', 'wb')
            image.write(text_binary)

            yet = ["AWS/RDS", "CPUUtilization", "DBInstanceIdentifier", run,
                   {"period": 300, "stat": "Average"}]
            # yet_data.append(yet)
            x = {
                "metrics": [yet],

                "legend": {
                    "position": "bottom"
                },
                "region": k,
                "liveData": False,
                "title": "CPUUtilization: Percent",
                "view": "timeSeries",
                "stacked": False
            }
            y = json.dumps(x)
            client = boto3.client('cloudwatch', region_name=k)
            response = client.get_metric_widget_image(MetricWidget=y)

            img = response['MetricWidgetImage']
            text_binary = img
            image = open('CPU' + k + '.png', 'wb')
            image.write(text_binary)

            net = ["AWS/RDS", "FreeStorageSpace", "DBInstanceIdentifier", run, {"period": 300, "stat": "Average"}]
            # net_data.append(net)

            m = {
                "sparkline": True,
                "metrics": [net],

                "legend": {
                    "position": "bottom"
                },
                "liveData": False,
                "title": "FreeStorageSpace: MB/Second",
                "view": "timeSeries",
                "stacked": False,
                "width": 600,
                "height": 400,
                "start": "-PT24H",
                "end": "P0D"
            }
            n = json.dumps(m)
            client = boto3.client('cloudwatch', region_name=k)
            response = client.get_metric_widget_image(MetricWidget=n)

            img = response['MetricWidgetImage']
            text_binary = img
            image = open('Freestg' + k + '.png', 'wb')
            image.write(text_binary)

            bet = ["AWS/RDS", "FreeableMemory", "DBInstanceIdentifier", run, {"period": 300, "stat": "Average"}]
            # bet_data.append(bet)

            a = {
                "metrics": [bet],

                "legend": {
                    "position": "bottom"
                },
                "liveData": False,
                "title": "FreeableMemory: MB",
                "view": "timeSeries",
                "stacked": False,
                "width": 600,
                "height": 400,
                "start": "-PT3H",
                "end": "P0D"
            }
            b = json.dumps(a)
            client = boto3.client('cloudwatch', region_name=k)
            response = client.get_metric_widget_image(MetricWidget=b)

            img = response['MetricWidgetImage']
            text_binary = img
            image = open('Freemem' + k + '.png', 'wb')
            image.write(text_binary)

            pet = ["AWS/RDS", "WriteIOPS", "DBInstanceIdentifier", run,
                   {"period": 300, "stat": "Average", "yAxis": "right"}]
            # pet_data.append(pet)
            o = {
                "metrics": [pet],
                "legend": {
                    "position": "bottom"
                },
                "liveData": False,
                "title": "WriteIOPS: Count/Second",
                "view": "timeSeries",
                "stacked": False,
                "width": 600,
                "height": 400,
                "start": "-PT3H",
                "end": "P0D"
            }
            p = json.dumps(o)
            client = boto3.client('cloudwatch', region_name=k)
            response = client.get_metric_widget_image(MetricWidget=p)

            img = response['MetricWidgetImage']
            text_binary = img
            image = open('writeiops' + k + '.png', 'wb')
            image.write(text_binary)

            det = ["AWS/RDS", "ReadIOPS", "DBInstanceIdentifier", run,
                   {"period": 300, "stat": "Average", "yAxis": "right"}]
            # det_data.append(det)
            c = {
                "metrics": [det],
                "legend": {
                    "position": "bottom"
                },
                "liveData": False,
                "title": "ReadIOPS: Count/Second",
                "view": "timeSeries",
                "stacked": False,
                "width": 600,
                "height": 400,
                "start": "-PT3H",
                "end": "P0D"
            }
            d = json.dumps(c)
            client = boto3.client('cloudwatch', region_name=k)
            response = client.get_metric_widget_image(MetricWidget=d)

            img = response['MetricWidgetImage']
            text_binary = img
            image = open('readiops' + k + '.png', 'wb')
            image.write(text_binary)

            sections = doc.sections
            for section in sections:
                section.top_margin = Mm(10)
                section.bottom_margin = Mm(10)
                section.left_margin = Mm(10)
                section.right_margin = Mm(10)

            para = doc.add_paragraph().add_run(' AWS Database Services')
            para.font.size = Pt(18)
            font = para.font
            font.color.rgb = RGBColor(0x42, 0x24, 0xE9)
            para = doc.add_paragraph().add_run(run)
            para.font.size = Pt(18)
            font = para.font
            font.color.rgb = RGBColor(255, 0, 0)

            paragraph = doc.add_paragraph()

            run = paragraph.add_run()
            run.add_picture('CPU' + k + '.png', width=Inches(2.15))
            run_2 = paragraph.add_run()
            run_2.add_picture('DatabaseConnections' + k + '.png', width=Inches(2.15))
            run_3 = paragraph.add_run()
            run_3.add_picture('Freestg' + k + '.png', width=Inches(2.15))
            run_4 = paragraph.add_run()
            run_4.add_picture('Freemem' + k + '.png', width=Inches(2.15))
            run_5 = paragraph.add_run()
            run_5.add_picture('writeiops' + k + '.png', width=Inches(2.15))
            run_6 = paragraph.add_run()
            run_6.add_picture('readiops' + k + '.png', width=Inches(2.15))

        doc.save('Rdscw1.docx')


ec2gr()
