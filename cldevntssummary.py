import boto3
import docx
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Inches, Pt, RGBColor, Mm

from toturial.SNS.linkfile import add_hyperlink

data = []


def eve():
    ec2client = boto3.client('ec2')
    regions = ec2client.describe_regions(
        AllRegions=False
    )
    for i in regions['Regions']:
        k = (i['RegionName'])

        client = boto3.client('events', region_name=k)
        response = client.list_rules()
        for rll in response['Rules']:
            rname = rll['Name']
            # print(rname)

            response1 = client.describe_rule(
                Name=rname)
            stat = response1['State']
            rll['State'] = stat
            try:
                shld = response1['ScheduleExpression']
                rll['EventPattern'] = shld
            except:
                rll['EventPattern'] = 'No ScheduleExpression'

            response2 = client.list_targets_by_rule(
                Rule=rname)
            for ttd in response2['Targets']:
                tid = ttd['Id']
                # print(tid)
                rll['Target'] = tid
                rll['Region'] = k
                data.append(rll)

            doc = docx.Document()

            sections = doc.sections
            for section in sections:
                section.top_margin = Mm(15)
                section.bottom_margin = Mm(15)
                section.left_margin = Mm(15)
                section.right_margin = Mm(15)
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(
                '                                                                                                                                                                           ')
            run.add_picture('cloudjournee.png', width=Inches(1.25))

            para = doc.add_paragraph().add_run('AWS Inventory - Current Consolidated Report')

            para.font.size = Pt(24)
            font = para.font
            font.color.rgb = RGBColor(0, 0, 255)
            p = doc.add_paragraph('The following line are the list of assets audited')
            add_hyperlink(p, 'Link to my site', "https://www.cloudjournee.com/")
            para = doc.add_paragraph().add_run('Cloud Watch Alarm Event Summary')
            para.font.size = Pt(18)
            font = para.font
            font.color.rgb = RGBColor(255, 0, 0)
            menutable = doc.add_table(rows=1, cols=5)
            menutable.style = 'Table Grid'

            shading_elm_1 = parse_xml(r'<w:shd {} w:fill="1F5C8B"/>'.format(nsdecls('w')))
            menutable.rows[0].cells[0]._tc.get_or_add_tcPr().append(shading_elm_1)
            shading_elm_1 = parse_xml(r'<w:shd {} w:fill="1F5C8B"/>'.format(nsdecls('w')))
            menutable.rows[0].cells[1]._tc.get_or_add_tcPr().append(shading_elm_1)
            shading_elm_1 = parse_xml(r'<w:shd {} w:fill="1F5C8B"/>'.format(nsdecls('w')))
            menutable.rows[0].cells[2]._tc.get_or_add_tcPr().append(shading_elm_1)
            shading_elm_1 = parse_xml(r'<w:shd {} w:fill="1F5C8B"/>'.format(nsdecls('w')))
            menutable.rows[0].cells[3]._tc.get_or_add_tcPr().append(shading_elm_1)
            shading_elm_1 = parse_xml(r'<w:shd {} w:fill="1F5C8B"/>'.format(nsdecls('w')))
            menutable.rows[0].cells[4]._tc.get_or_add_tcPr().append(shading_elm_1)
            hdr_cells = menutable.rows[0].cells
            hdr_cells[0].text = 'Name'
            hdr_cells[1].text = 'State'
            hdr_cells[2].text = 'EventPattern'
            hdr_cells[3].text = 'Region'
            hdr_cells[4].text = 'Target'

            for i in data:
                # print(i)
                row_cells = menutable.add_row().cells
                row_cells[0].text = i['Name']
                row_cells[1].text = i['State']
                row_cells[2].text = i['EventPattern']
                row_cells[3].text = i['Region']
                row_cells[4].text = i['Target']

            for row in menutable.rows:
                for cell in row.cells:
                    paragraphs = cell.paragraphs
                    for paragraph in paragraphs:
                        for run in paragraph.runs:
                            font = run.font
                            font.size = Pt(10)

            doc.save('cldevnt.docx')


eve()
