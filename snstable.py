import boto3
import docx
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Mm, Inches

from toturial.SNS.linkfile import add_hyperlink
from toturial.SNS.serializers import sn

data = []


def snss():
    ec2client = boto3.client('ec2')
    regions = ec2client.describe_regions(
        AllRegions=False
    )
    for i in regions['Regions']:
        k = (i['RegionName'])
        s_client = boto3.client('sns', region_name=k)
        response = s_client.list_subscriptions()
        for sss in response['Subscriptions']:
            sun = sss['SubscriptionArn']
            sss['region'] = k
            data.append(sss)
    #         serial = sn(data=sss)
    #         if serial.is_valid():
    #             data.append(serial.data)
    #         else:
    #             print(serial.errors)
    # return data
    doc = docx.Document()

    sections = doc.sections
    for section in sections:
        section.top_margin = Mm(10)
        section.bottom_margin = Mm(10)
        section.left_margin = Mm(10)
        section.right_margin = Mm(10)
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(
        '                                                                                                                                                                           ')
    picture = run.add_picture('cloudjournee.png', width=Inches(1.25))

    para = doc.add_paragraph().add_run('AWS Inventory - Current Consolidated Report')

    para.font.size = Pt(24)
    font = para.font
    font.color.rgb = RGBColor(0, 0, 255)
    p = doc.add_paragraph('The following line are the list of assets audited')

    add_hyperlink(p,   'Link to my site', "https://www.cloudjournee.com/")
    para = doc.add_paragraph().add_run('SNS Summary')
    para.font.size = Pt(18)
    font = para.font
    font.color.rgb = RGBColor(255, 0, 0)
    menutable = doc.add_table(rows=1, cols=5)
    menutable.style = 'Colorful List'
    hdr_cells = menutable.rows[0].cells
    hdr_cells[0].text = 'TopicArn'
    hdr_cells[1].text = 'region'
    hdr_cells[2].text = 'SubscriptionArn'
    hdr_cells[3].text = 'Protocol'
    hdr_cells[4].text = 'Endpoint'
    for i in data:
        print(i)
        row_cells = menutable.add_row().cells
        row_cells[0].text = i['TopicArn']
        row_cells[1].text = i['region']
        row_cells[2].text = i['SubscriptionArn']
        row_cells[3].text = i['Protocol']
        row_cells[4].text = i['Endpoint']

        for row in menutable.rows:
            for cell in row.cells:
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    for run in paragraph.runs:
                        font = run.font
                        font.size = Pt(7)

    cell_xml_element = menutable.rows[2].cells[0]._tc
    table_cell_properties = cell_xml_element.get_or_add_tcPr()
    shade_obj = OxmlElement('w:shd')
    shade_obj.set(qn('w:fill'), "#00ff00")
    table_cell_properties.append(shade_obj)

    cell_xml_element1 = menutable.rows[5].cells[2]._tc
    table_cell_properties = cell_xml_element1.get_or_add_tcPr()
    shade_obj = OxmlElement('w:shd')
    shade_obj.set(qn('w:fill'), "#00ff00")
    table_cell_properties.append(shade_obj)

    doc.save('sns.docx')


snss()
