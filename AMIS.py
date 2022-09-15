import boto3
import docx
from docx import Document
from docx.shared import Inches, Mm, Pt, RGBColor

from toturial.ec2.ec2link import add_hyperlink

data = []


def ec2data():
    ec2client = boto3.client('ec2')
    regions = ec2client.describe_regions(
        AllRegions=False
    )
    for i in regions['Regions']:
        k = (i['RegionName'])
        ec2_client = boto3.client('ec2', region_name=k)

        response1 = ec2_client.describe_images(Owners=[
            'self'
        ])
        for sna in response1['Images']:
            # print(k, sna)
            nname = sna['Name']
            img = sna['ImageId']

            # print(img)
            for bd in sna['BlockDeviceMappings']:

                try:
                    bd['SnapshotId'] = bd['Ebs']['SnapshotId']
                    bbb = bd['SnapshotId']
                    sna['SnapshotId'] = bbb

                    # print(nname, bbb, k)
                except:
                    print('  ')
            response = ec2_client.describe_instances(Filters=[
                {
                    'Name': 'image-id',
                    'Values': [
                        img,
                    ]
                },
            ], )
            sna['RegionName'] = k
            for ec2da in response['Reservations']:
                # print(ec2da)
                for instances in ec2da['Instances']:
                    insd = instances['InstanceId']
                    # print(insd)
                    sna['InstanceId'] = insd
            data.append(sna)
    doc = docx.Document()

    sections = doc.sections
    for section in sections:
        section.top_margin = Mm(15)
        section.bottom_margin = Mm(15)
        section.left_margin = Mm(15)
        section.right_margin = Mm(15)
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(
        '                                                                                                                                                                           ')
    run.add_picture('cloudjournee.png', width=Inches(1.25))

    para = doc.add_paragraph().add_run('AWS Inventory - Current Consolidated Report')

    para.font.size = Pt(24)
    font = para.font
    font.color.rgb = RGBColor(0, 0, 255)
    p = doc.add_paragraph('The following line are the list of assets audited')
    add_hyperlink(p, 'Link to my site', "https://www.cloudjournee.com/")
    para = doc.add_paragraph().add_run('AMI Summary')
    para.font.size = Pt(18)
    font = para.font
    font.color.rgb = RGBColor(255, 0, 0)
    menutable = doc.add_table(rows=1, cols=5)
    menutable.style = 'Colorful List'
    hdr_cells = menutable.rows[0].cells
    hdr_cells[0].text = 'ImageName'
    hdr_cells[1].text = 'ImageId'
    hdr_cells[2].text = 'SnapshotId'
    hdr_cells[3].text = 'InstanceId'
    hdr_cells[4].text = 'region'

    for i in data:
        # print(i)
        row_cells = menutable.add_row().cells
        try:
            row_cells[0].text = i['Name']
            row_cells[1].text = i['ImageId']
            row_cells[2].text = i['SnapshotId']
            row_cells[4].text = i['RegionName']
            row_cells[3].text = i['InstanceId']

        except:
            row_cells[3].text = ['  -   ']

    for row in menutable.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size = Pt(10)
    doc.save('AMI.docx')


ec2data()
