from datetime import datetime

import boto3
import docx
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Pt, RGBColor, Inches, Mm

from toturial.SNS.linkfile import add_hyperlink

data = []

def alm():
    ec2client = boto3.client('ec2')
    regions = ec2client.describe_regions(
        AllRegions=False
    )
    for i in regions['Regions']:
        k = (i['RegionName'])
        client = boto3.client('cloudwatch', region_name=k)
        response = client.describe_alarms()
        alc = response['CompositeAlarms']
        for ala in response['MetricAlarms']:
            aln = ala['AlarmName']
            alup = ala['AlarmConfigurationUpdatedTimestamp']
            statv = ala['StateValue']
            mtn = ala['MetricName']
            nas = ala['Namespace']

            for alac in ala['AlarmActions']:
                alamact = alac
                ala['AlarmActions'] = alamact
            for dem in ala['Dimensions']:
                dname = dem['Name']
                ala['DimensionName'] = dname
            data.append(ala)

        doc = docx.Document()

        sections = doc.sections
        for section in sections:
            section.top_margin = Mm(15)
            section.bottom_margin = Mm(15)
            section.left_margin = Mm(15)
            section.right_margin = Mm(15)
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(
            '                                                                                                                                                                           ')
        run.add_picture('cloudjournee.png', width=Inches(1.25))

        para = doc.add_paragraph().add_run('AWS Inventory - Current Consolidated Report')

        para.font.size = Pt(24)
        font = para.font
        font.color.rgb = RGBColor(0, 0, 255)
        p = doc.add_paragraph('The following line are the list of assets audited')
        add_hyperlink(p, 'Link to my site', "https://www.cloudjournee.com/")
        para = doc.add_paragraph().add_run('Cloud Watch Alarm Summary')
        para.font.size = Pt(18)
        font = para.font
        font.color.rgb = RGBColor(255, 0, 0)
        menutable = doc.add_table(rows=1, cols=7)
        menutable.style = 'Table Grid'

        shading_elm_1 = parse_xml(r'<w:shd {} w:fill="1F5C8B"/>'.format(nsdecls('w')))
        menutable.rows[0].cells[0]._tc.get_or_add_tcPr().append(shading_elm_1)
        shading_elm_1 = parse_xml(r'<w:shd {} w:fill="1F5C8B"/>'.format(nsdecls('w')))
        menutable.rows[0].cells[1]._tc.get_or_add_tcPr().append(shading_elm_1)
        shading_elm_1 = parse_xml(r'<w:shd {} w:fill="1F5C8B"/>'.format(nsdecls('w')))
        menutable.rows[0].cells[2]._tc.get_or_add_tcPr().append(shading_elm_1)
        shading_elm_1 = parse_xml(r'<w:shd {} w:fill="1F5C8B"/>'.format(nsdecls('w')))
        menutable.rows[0].cells[3]._tc.get_or_add_tcPr().append(shading_elm_1)
        shading_elm_1 = parse_xml(r'<w:shd {} w:fill="1F5C8B"/>'.format(nsdecls('w')))
        menutable.rows[0].cells[4]._tc.get_or_add_tcPr().append(shading_elm_1)
        shading_elm_1 = parse_xml(r'<w:shd {} w:fill="1F5C8B"/>'.format(nsdecls('w')))
        menutable.rows[0].cells[5]._tc.get_or_add_tcPr().append(shading_elm_1)
        shading_elm_1 = parse_xml(r'<w:shd {} w:fill="1F5C8B"/>'.format(nsdecls('w')))
        menutable.rows[0].cells[6]._tc.get_or_add_tcPr().append(shading_elm_1)

        hdr_cells = menutable.rows[0].cells
        hdr_cells[0].text = 'AlarmName'
        hdr_cells[1].text = 'Alarm Updated'
        hdr_cells[2].text = 'AlarmActions'
        hdr_cells[3].text = 'StateValue'
        hdr_cells[4].text = 'MetricName'
        hdr_cells[5].text = 'Namespace'
        hdr_cells[6].text = 'Dimension Name'
        for i in data:
            # print(i)
            row_cells = menutable.add_row().cells
            row_cells[0].text = i['AlarmName']
            row_cells[1].text = str(i['AlarmConfigurationUpdatedTimestamp'])
            row_cells[2].text = i['AlarmActions']
            row_cells[3].text = i['StateValue']
            row_cells[4].text = i['MetricName']
            row_cells[5].text = i['Namespace']
            row_cells[6].text = i['DimensionName']

        for row in menutable.rows:
            for cell in row.cells:
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    for run in paragraph.runs:
                        font = run.font
                        font.size = Pt(10)

        doc.save('alarmsum.docx')


alm()
