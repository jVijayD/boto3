import boto3
import docx
from docx.opc.oxml import qn
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Mm, Inches, RGBColor, Pt

from toturial.det.serializers import s33

data = []


def sss():
    s3_client = boto3.client('s3')
    response = s3_client.list_buckets()
    for sb in response['Buckets']:
        sname = sb['Name']
        response1 = s3_client.get_bucket_logging(Bucket=sname)
        print(response1)
        try:
            log = response1['LoggingEnabled']
            print(log)
        except:

            # data.append(data)

            doc = docx.Document()

            sections = doc.sections
            for section in sections:
                section.top_margin = Mm(15)
                section.bottom_margin = Mm(15)
                section.left_margin = Mm(15)
                section.right_margin = Mm(15)
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(
                '                                                                                                                                                                           ')
            run.add_picture('cloudjournee.png', width=Inches(1.25))

            para = doc.add_paragraph().add_run('AWS Inventory - Current Consolidated Report')

            para.font.size = Pt(24)
            font = para.font
            font.color.rgb = RGBColor(0, 0, 255)
            menutable = doc.add_table(rows=2, cols=4)
            menutable.style = 'Table Grid'
            shading_elm_1 = parse_xml(r'<w:shd {} w:fill="1F5C8B"/>'.format(nsdecls('w')))
            menutable.rows[0].cells[0]._tc.get_or_add_tcPr().append(shading_elm_1)
            shading_elm_1 = parse_xml(r'<w:shd {} w:fill="1F5C8B"/>'.format(nsdecls('w')))
            menutable.rows[0].cells[1]._tc.get_or_add_tcPr().append(shading_elm_1)
            shading_elm_1 = parse_xml(r'<w:shd {} w:fill="1F5C8B"/>'.format(nsdecls('w')))
            menutable.rows[0].cells[2]._tc.get_or_add_tcPr().append(shading_elm_1)
            shading_elm_1 = parse_xml(r'<w:shd {} w:fill="1F5C8B"/>'.format(nsdecls('w')))
            menutable.rows[0].cells[3]._tc.get_or_add_tcPr().append(shading_elm_1)
            shading_elm_1 = parse_xml(r'<w:shd {} w:fill="#FFFF00"/>'.format(nsdecls('w')))
            menutable.rows[1].cells[2]._tc.get_or_add_tcPr().append(shading_elm_1)
            hdr_cells = menutable.rows[0].cells
            hdr_cells[0].text = 'AssetType'
            hdr_cells[1].text = 'Risk/Potential Gap'
            hdr_cells[2].text = 'Severity'
            hdr_cells[3].text = 'Description'

            row_cells = menutable.rows[1].cells
            row_cells[0].text = ['s3']
            row_cells[1].text = ['AWS s3 Bucket Logging level is disabled']
            row_cells[2].text = ['Medium']
            row_cells[3].text = [
                'If loging is not enabled, s3 bucket Owners will not be able to determine who is accessing']

            for row in menutable.rows:
                for cell in row.cells:
                    paragraphs = cell.paragraphs
                    for paragraph in paragraphs:
                        for run in paragraph.runs:
                            font = run.font
                            font.size = Pt(10)
            widths = (Inches(0.5), Inches(2), Inches(0.5), Inches(4))
            for row in menutable.rows:
                for idx, width in enumerate(widths):
                    row.cells[idx].width = width
            doc.save('s3Buckets.docx')


sss()
