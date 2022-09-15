import boto3
import docx
from docx.shared import Mm, Pt, RGBColor, Inches

from toturial.RDS.serilizers import rd

data = []


def rdds():
    ec2client = boto3.client('ec2')
    regions = ec2client.describe_regions(
        AllRegions=False
    )
    for i in regions['Regions']:
        k = (i['RegionName'])
    rdclient = boto3.client('rds', region_name=k)
    response = rdclient.describe_db_instances()
    for rrr in response['DBInstances']:
        run = rrr['DBInstanceIdentifier']
        rrn = rrr['DBSubnetGroup']['VpcId']
        rrr['VpcId'] = rrn
        rrr['region'] = k
        data.append(rrr)

    doc = docx.Document()

    sections = doc.sections
    for section in sections:
        section.top_margin = Mm(15)
        section.bottom_margin = Mm(15)
        section.left_margin = Mm(15)
        section.right_margin = Mm(15)
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(
        '                                                                                                                                                                           ')
    picture = run.add_picture('cloudjournee.png', width=Inches(1.25))

    para = doc.add_paragraph().add_run('AWS Inventory - Current Consolidated Report')

    para.font.size = Pt(24)
    font = para.font
    font.color.rgb = RGBColor(0, 0, 255)
    p = doc.add_paragraph('The following line are the list of assets audited')
    para = doc.add_paragraph().add_run('RDS Summary')
    para.font.size = Pt(18)
    font = para.font
    font.color.rgb = RGBColor(255, 0, 0)
    menutable = doc.add_table(rows=1, cols=11)
    menutable.style = 'Colorful List'
    hdr_cells = menutable.rows[0].cells
    hdr_cells[0].text = 'DBInstanceIdentifier'
    hdr_cells[1].text = 'DBInstanceClass'
    hdr_cells[2].text = 'Engine'
    hdr_cells[3].text = 'StorageType'
    hdr_cells[4].text = 'InstanceCreateTime'
    hdr_cells[5].text = 'DBInstanceStatus'
    hdr_cells[6].text = 'VpcId'
    hdr_cells[7].text = 'MultiAZ'
    hdr_cells[8].text = 'BackupRetentionPeriod'
    hdr_cells[9].text = 'AvailabilityZone'
    hdr_cells[10].text = 'region'
    for i in data:
        print(i)
        row_cells = menutable.add_row().cells
        row_cells[0].text = i['DBInstanceIdentifier']
        row_cells[1].text = i['DBInstanceClass']
        row_cells[2].text = i['Engine']
        row_cells[3].text = i['StorageType']
        row_cells[4].text = str(i['InstanceCreateTime'])
        row_cells[5].text = i['DBInstanceStatus']
        row_cells[6].text = i['VpcId']
        row_cells[7].text = str(i['MultiAZ'])
        row_cells[8].text = str(i['BackupRetentionPeriod'])
        row_cells[9].text = i['AvailabilityZone']
        row_cells[10].text = i['region']

        for row in menutable.rows:
            for cell in row.cells:
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    for run in paragraph.runs:
                        font = run.font
                        font.size = Pt(8)

    doc.save('rds.docx')


rdds()
