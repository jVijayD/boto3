from datetime import date

import boto3
import docx
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Mm, Inches, RGBColor, Pt


def iac():
    iamclient = boto3.client('iam')
    response = iamclient.list_users()
    for udd in response['Users']:
        usr = udd['UserName']
        response1 = iamclient.list_access_keys(UserName=usr)
        for cai in response1['AccessKeyMetadata']:
            print(cai)
            actv = cai['Status']
            accesskeydate = response1['AccessKeyMetadata'][0]['CreateDate'].date()
            currentdate = date.today()
            active_days = currentdate - accesskeydate

            if active_days.days > 90:
                print(' IAM user has Programmatic access beyond 90 days time period.', actv, usr)

                doc = docx.Document()

                sections = doc.sections
                for section in sections:
                    section.top_margin = Mm(10)
                    section.bottom_margin = Mm(10)
                    section.left_margin = Mm(10)
                    section.right_margin = Mm(10)
                paragraph = doc.add_paragraph()
                run = paragraph.add_run(
                    '                                                                                                                                                                           ')
                run.add_picture('cloudjournee.png', width=Inches(1.25))

                para = doc.add_paragraph().add_run('AWS Inventory - Current Consolidated Report')

                para.font.size = Pt(24)
                font = para.font
                font.color.rgb = RGBColor(0, 0, 255)
                menutable = doc.add_table(rows=1, cols=4)

                menutable.style = 'Table Grid'

                shading_elm_1 = parse_xml(r'<w:shd {} w:fill="1F5C8B"/>'.format(nsdecls('w')))
                menutable.rows[0].cells[0]._tc.get_or_add_tcPr().append(shading_elm_1)
                shading_elm_1 = parse_xml(r'<w:shd {} w:fill="1F5C8B"/>'.format(nsdecls('w')))
                menutable.rows[0].cells[1]._tc.get_or_add_tcPr().append(shading_elm_1)
                shading_elm_1 = parse_xml(r'<w:shd {} w:fill="1F5C8B"/>'.format(nsdecls('w')))
                menutable.rows[0].cells[2]._tc.get_or_add_tcPr().append(shading_elm_1)
                shading_elm_1 = parse_xml(r'<w:shd {} w:fill="1F5C8B"/>'.format(nsdecls('w')))
                menutable.rows[0].cells[3]._tc.get_or_add_tcPr().append(shading_elm_1)
                shading_elm_1 = parse_xml(r'<w:shd {} w:fill="#FF0000"/>'.format(nsdecls('w')))

                hdr_cells = menutable.rows[0].cells
                hdr_cells[0].text = 'AssetType'
                hdr_cells[1].text = 'Risk/Potential Gap'
                hdr_cells[2].text = 'Severity'
                hdr_cells[3].text = 'Description'

                row_cells = menutable.add_row().cells
                # row_cells = menutable.rows[1].cells
                row_cells[0].text = ['IAM']
                row_cells[1].text = ['Checks if an  IAM User has Programmatic access with active access key']
                row_cells[2].text = ['High']
                row_cells[3].text = [
                    'It is recommended that IAM User shouldn’t have the Programmatic access beyond certain period']
                menutable.rows[1].cells[2]._tc.get_or_add_tcPr().append(shading_elm_1)
                for row in menutable.rows:
                    for cell in row.cells:
                        paragraphs = cell.paragraphs
                        for paragraph in paragraphs:
                            for run in paragraph.runs:
                                font = run.font
                                font.size = Pt(10)
                widths = (Inches(0.5), Inches(2), Inches(0.5), Inches(4))
                for row in menutable.rows:
                    for idx, width in enumerate(widths):
                        row.cells[idx].width = width
                doc.save('prgmacces.docx')
            else:
                print('Users has Active Key', actv, usr)


iac()
