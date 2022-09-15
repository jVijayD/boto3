import json
from datetime import datetime

import boto3
import docx
from docx.shared import Pt, RGBColor, Inches, Mm

from toturial.SNS.linkfile import add_hyperlink

data = []
sub = []


def sqs():
    ec2client = boto3.client('ec2')
    regions = ec2client.describe_regions(
        AllRegions=False
    )
    for i in regions['Regions']:
        k = (i['RegionName'])

        sqs_client = boto3.client("sqs", region_name=k)
        response = sqs_client.list_queues()
        # print(response)
        try:
            qurl = response['QueueUrls']
            for end in qurl:
                dub = {}

                # print(end)
                response1 = sqs_client.get_queue_attributes(
                    QueueUrl=end,
                    AttributeNames=['All'])
                # print(response1)
                att = response1['Attributes']
                qrn = att['QueueArn']
                cdt = att['CreatedTimestamp']
                ldt = att['LastModifiedTimestamp']
                ts = int(att['CreatedTimestamp'])
                crtm = (datetime.utcfromtimestamp(ts).strftime('%Y-%m-%d %H:%M:%S'))
                ts1 = int(att['LastModifiedTimestamp'])
                lmtm = (datetime.utcfromtimestamp(ts).strftime('%Y-%m-%d %H:%M:%S'))
                att['CreatedTimestamp'] = crtm
                att['LastModifiedTimestamp'] = lmtm
                try:
                    ppc = att['Policy']
                    ppd = json.loads(ppc)
                    ppe = ppd['Id']
                    att['PolicyId'] = ppe
                except:
                    ppe = 'Deafult policy'
                    att['PolicyId'] = ppe
                # print(att, qrn)
                try:
                    if att['FifoQueue']:
                        # print('fifo', k)
                        dub['QueueType'] = 'FIFO'

                except:
                    # print('Standard', k)
                    dub['QueueType'] = 'Standard'

                client = boto3.client("sts")
                acid = (client.get_caller_identity()["Account"])
                qname = (qrn.split(acid + ':')[1])
                dub['QueueName'] = qname
                # dub['PolicyId'] = ppe
                dub['region'] = k
                dub['att'] = att
                sub.append(dub)
                # sub.append(att)
                # print(data)
        except:
            print('no Queue', k)

    # print(sub)
    doc = docx.Document()

    sections = doc.sections
    for section in sections:
        section.top_margin = Mm(15)
        section.bottom_margin = Mm(15)
        section.left_margin = Mm(15)
        section.right_margin = Mm(15)
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(
        '                                                                                                                                                                           ')
    run.add_picture('cloudjournee.png', width=Inches(1.25))

    para = doc.add_paragraph().add_run('AWS Inventory - Current Consolidated Report')

    para.font.size = Pt(24)
    font = para.font
    font.color.rgb = RGBColor(0, 0, 255)
    p = doc.add_paragraph('The following line are the list of assets audited')
    add_hyperlink(p, 'Link to my site', "https://www.cloudjournee.com/")
    para = doc.add_paragraph().add_run('SQS Summary')
    para.font.size = Pt(18)
    font = para.font
    font.color.rgb = RGBColor(255, 0, 0)
    menutable = doc.add_table(rows=1, cols=13)
    menutable.style = 'Colorful List'
    hdr_cells = menutable.rows[0].cells
    hdr_cells[0].text = 'QueueName'
    hdr_cells[1].text = 'QueueType'
    hdr_cells[2].text = 'Created'
    hdr_cells[3].text = 'LastModified'
    hdr_cells[4].text = 'region'
    hdr_cells[5].text = 'DelaySeconds'
    hdr_cells[6].text = 'Messages'
    hdr_cells[7].text = 'Message Not Visible'
    hdr_cells[8].text = 'Messages Delayed'
    hdr_cells[9].text = 'VisibilityTimeout'
    hdr_cells[10].text = 'MaximumMessageSize'
    hdr_cells[11].text = 'MessageRetentionPeriod'
    hdr_cells[12].text = 'Policy'
    for i in sub:
        row_cells = menutable.add_row().cells
        row_cells[0].text = i['QueueName']
        row_cells[1].text = i['QueueType']
        row_cells[2].text = i['att']['CreatedTimestamp']
        row_cells[3].text = i['att']['LastModifiedTimestamp']
        row_cells[4].text = i['region']
        row_cells[5].text = i['att']['DelaySeconds']
        row_cells[6].text = i['att']['ApproximateNumberOfMessages']
        row_cells[7].text = i['att']['ApproximateNumberOfMessagesNotVisible']
        row_cells[8].text = i['att']['ApproximateNumberOfMessagesDelayed']
        row_cells[9].text = i['att']['VisibilityTimeout']
        row_cells[10].text = i['att']['MaximumMessageSize']
        row_cells[11].text = i['att']['MessageRetentionPeriod']
        row_cells[12].text = i['att']['PolicyId']

    for row in menutable.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size = Pt(7)
    doc.save('SQS.docx')


sqs()
