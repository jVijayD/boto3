import boto3
import docx
from docx.shared import Mm, Inches, Pt, RGBColor

from toturial.ebs.serializers import eb
from toturial.ec2.ec2link import add_hyperlink

data = []


def ebbs():
    ec2client = boto3.client('ec2')
    regions = ec2client.describe_regions(
        AllRegions=False
    )
    for i in regions['Regions']:
        k = (i['RegionName'])
        ebclient = boto3.client('ec2', region_name=k)
        response = ebclient.describe_volumes()
        for eee in response['Volumes']:
            for ins in eee['Attachments']:
                bun = ins['InstanceId']
                eee['InstanceId'] = bun
                eee['region'] = k
                data.append(eee)
    #             serial = eb(data=eee)
    #             if serial.is_valid():
    #                 data.append(serial.data)
    #             else:
    #                 print(serial.errors)
    #
    # return data

    doc = docx.Document()

    sections = doc.sections
    for section in sections:
        section.top_margin = Mm(15)
        section.bottom_margin = Mm(15)
        section.left_margin = Mm(15)
        section.right_margin = Mm(15)
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(
        '                                                                                                                                                                           ')
    picture = run.add_picture('cloudjournee.png', width=Inches(1.25))

    para = doc.add_paragraph().add_run('AWS Inventory - Current Consolidated Report')

    para.font.size = Pt(24)
    font = para.font
    font.color.rgb = RGBColor(0, 0, 255)
    p = doc.add_paragraph('The following line are the list of assets audited')
    add_hyperlink(p, 'Link to my site', "https://www.cloudjournee.com/")
    para = doc.add_paragraph().add_run('EBS Summary')
    para.font.size = Pt(18)
    font = para.font
    font.color.rgb = RGBColor(255, 0, 0)
    menutable = doc.add_table(rows=1, cols=8)
    menutable.style = 'Colorful List'
    hdr_cells = menutable.rows[0].cells
    hdr_cells[0].text = 'VolumeId'
    hdr_cells[1].text = 'VolumeType'
    hdr_cells[2].text = 'Size'
    hdr_cells[3].text = 'CreateTime'
    hdr_cells[4].text = 'AvailabilityZone'
    hdr_cells[5].text = 'region'
    hdr_cells[6].text = 'State'
    hdr_cells[7].text = 'InstanceId'
    for i in data:
        print(i)
        row_cells = menutable.add_row().cells
        row_cells[0].text = i['VolumeId']
        row_cells[1].text = i['VolumeType']
        row_cells[2].text = str(i['Size'])
        row_cells[3].text = str(i['CreateTime'])
        row_cells[4].text = i['AvailabilityZone']
        row_cells[5].text = i['region']
        row_cells[6].text = i['State']
        row_cells[7].text = i['InstanceId']

        for row in menutable.rows:
            for cell in row.cells:
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    for run in paragraph.runs:
                        font = run.font
                        font.size = Pt(8)

    doc.save('ebs.docx')


ebbs()
