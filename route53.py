import boto3
import docx
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Pt, RGBColor, Inches, Mm

from toturial.SNS.linkfile import add_hyperlink

data = []


def rot():
    client = boto3.client('route53')
    response = client.list_hosted_zones()
    # print(response)
    for hz in response['HostedZones']:
        id = hz['Id']
        dname = hz['Name']
        count = hz['ResourceRecordSetCount']
        config = hz['Config']['PrivateZone']
        hz['PrivateZone'] = config

        spid = id.split('/hostedzone/')
        hz['Id'] = spid
        # print(count)
        # response1 = client.list_resource_record_sets(
        #     HostedZoneId=id)
        # print()
        response2 = client.list_health_checks()
        for hc in response2['HealthChecks']:
            hid = hc['Id']
            # response3 = client.get_health_check_status(
            #     HealthCheckId=hid)
            print(response)
        data.append(hz)

        doc = docx.Document()

        sections = doc.sections
        for section in sections:
            section.top_margin = Mm(15)
            section.bottom_margin = Mm(15)
            section.left_margin = Mm(15)
            section.right_margin = Mm(15)
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(
            '                                                                                                                                                                           ')
        run.add_picture('cloudjournee.png', width=Inches(1.25))

        para = doc.add_paragraph().add_run('AWS Inventory - Current Consolidated Report')

        para.font.size = Pt(24)
        font = para.font
        font.color.rgb = RGBColor(0, 0, 255)
        p = doc.add_paragraph('The following line are the list of assets audited')
        add_hyperlink(p, 'Link to my site', "https://www.cloudjournee.com/")
        para = doc.add_paragraph().add_run('Route53 Summary')
        para.font.size = Pt(18)
        font = para.font
        font.color.rgb = RGBColor(255, 0, 0)
        menutable = doc.add_table(rows=1, cols=6)
        menutable.style = 'Table Grid'

        shading_elm_1 = parse_xml(r'<w:shd {} w:fill="1F5C8B"/>'.format(nsdecls('w')))
        menutable.rows[0].cells[0]._tc.get_or_add_tcPr().append(shading_elm_1)
        shading_elm_1 = parse_xml(r'<w:shd {} w:fill="1F5C8B"/>'.format(nsdecls('w')))
        menutable.rows[0].cells[1]._tc.get_or_add_tcPr().append(shading_elm_1)
        shading_elm_1 = parse_xml(r'<w:shd {} w:fill="1F5C8B"/>'.format(nsdecls('w')))
        menutable.rows[0].cells[2]._tc.get_or_add_tcPr().append(shading_elm_1)
        shading_elm_1 = parse_xml(r'<w:shd {} w:fill="1F5C8B"/>'.format(nsdecls('w')))
        menutable.rows[0].cells[3]._tc.get_or_add_tcPr().append(shading_elm_1)
        shading_elm_1 = parse_xml(r'<w:shd {} w:fill="1F5C8B"/>'.format(nsdecls('w')))
        menutable.rows[0].cells[4]._tc.get_or_add_tcPr().append(shading_elm_1)


        hdr_cells = menutable.rows[0].cells
        hdr_cells[0].text = 'Sl No'
        hdr_cells[1].text = 'Id'
        hdr_cells[2].text = 'Name'
        hdr_cells[3].text = 'Private Zone'
        hdr_cells[4].text = 'Resource RecordSet Count'
        hdr_cells[5].text = 'Status'
        q = 1
        for i in data:
            # print(i)
            row_cells = menutable.add_row().cells
            row_cells[0].text = str(q)
            row_cells[1].text = i['Id']
            row_cells[2].text = i['Name']
            row_cells[3].text = str(i['PrivateZone'])
            row_cells[4].text = str(i['ResourceRecordSetCount'])
            # row_cells[4].text = i['Status']
            q = q + 1
            for row in menutable.rows:
                for cell in row.cells:
                    paragraphs = cell.paragraphs
                    for paragraph in paragraphs:
                        for run in paragraph.runs:
                            font = run.font
                            font.size = Pt(10)

        doc.save('route53.docx')


rot()
