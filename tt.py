import docx
from docx import Document
from docx.shared import Inches, Mm

doc = docx.Document()

sections = doc.sections
for section in sections:
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)
    section.left_margin = Mm(25)
    section.right_margin = Mm(10)
paragraph = doc.add_paragraph()
run = paragraph.add_run(
    '                                                                                                             ')
picture = run.add_picture('cloudjournee.png', width=Inches(1.80))

doc.add_heading('AWS Inventory - Current Consolidated Report', 0)

p = doc.add_paragraph('The following line are the list of assets audited')
doc.add_heading('AMI Summary ', level=1)

records = (
    (3, '101', 'Spam', 'image', 'Snap', 'did'),
    (7, '422', 'Eggs', 't2.micro', 'iam', 'didnot'),
    (4, '631', 'Spam', 't3.micro', 'gp2', 'done')
)

menutable = doc.add_table(rows=1, cols=6)
menutable.style = 'Colourful List'
hdr_cells = menutable.rows[0].cells
hdr_cells[0].text = 'S_NO'
hdr_cells[1].text = 'Image Name'
hdr_cells[2].text = 'ImageID'
hdr_cells[3].text = 'SnapshotID'
hdr_cells[4].text = 'InstanceID'
hdr_cells[5].text = 'Region'
for S_NO, ImageName, ImageID, SnapshotID, InstanceID, Region in records:
    row_cells = menutable.add_row().cells
    row_cells[0].text = str(S_NO)
    row_cells[1].text = ImageName
    row_cells[2].text = ImageID
    row_cells[3].text = SnapshotID
    row_cells[4].text = InstanceID
    row_cells[5].text = Region
doc.add_page_break()

doc.save('demo.docx')
