import json

import boto3
import docx
from docx.shared import Pt, RGBColor, Inches, Mm

from toturial.lambd1.serializers import lamb

data = []


def lambdas():
    ec2client = boto3.client('ec2')
    regions = ec2client.describe_regions(
        AllRegions=False
    )
    for i in regions['Regions']:
        k = (i['RegionName'])
        l_client = boto3.client('lambda', region_name=k)
        response = l_client.list_functions()
        for lll in response['Functions']:
            fun = lll['FunctionName']
            lll['region'] = k
            data.append(lll)
            # print(fun)
        #     serial = lamb(data=lll)
        #     if serial.is_valid():
        #         data.append(serial.data)
        #     else:
        #         print(serial.errors)
        #
        # return data

    doc = docx.Document()

    sections = doc.sections
    for section in sections:
        section.top_margin = Mm(15)
        section.bottom_margin = Mm(15)
        section.left_margin = Mm(15)
        section.right_margin = Mm(15)
    paragraph = doc.add_paragraph()
    run = paragraph.add_run('                                                                                                 ')
    run.add_picture('cloudjournee.png', width=Inches(1.25))

    para = doc.add_paragraph().add_run('AWS Inventory - Current Consolidated Report')

    para.font.size = Pt(24)
    font = para.font
    font.color.rgb = RGBColor(0, 0, 255)
    p = doc.add_paragraph('The following line are the list of assets audited')
    para = doc.add_paragraph().add_run('Lambda Summary')
    para.font.size = Pt(18)
    font = para.font
    font.color.rgb = RGBColor(255, 0, 0)
    menutable = doc.add_table(rows=1, cols=4)
    menutable.style = 'Colorful List'
    hdr_cells = menutable.rows[0].cells
    hdr_cells[0].text = 'FunctionName'
    hdr_cells[1].text = 'Runtime'
    hdr_cells[2].text = 'Role'
    hdr_cells[3].text = 'region'
    for i in data:
        # print(i)
        row_cells = menutable.add_row().cells
        row_cells[0].text = i['FunctionName']
        row_cells[1].text = str(i['Runtime'])
        row_cells[2].text = i['Role']
        row_cells[3].text = i['region']

        for row in menutable.rows:
            for cell in row.cells:
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    for run in paragraph.runs:
                        font = run.font
                        font.size = Pt(10)

    doc.save('lambdaa.docx')


lambdas()
