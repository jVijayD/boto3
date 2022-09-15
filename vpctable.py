import boto3
import docx
from docx.shared import Mm, Inches, Pt, RGBColor

from toturial.ec2.ec2link import add_hyperlink
from toturial.vpc.serializer import vpc

data = []


def vpcdd():
    ec2client = boto3.client('ec2')
    regions = ec2client.describe_regions(
        AllRegions=False
    )
    for i in regions['Regions']:
        k = (i['RegionName'])
        ec2_client = boto3.client('ec2', region_name=k)
        response = ec2_client.describe_vpcs()
        for vpcd1 in response['Vpcs']:
            VpcID = vpcd1['VpcId']
            if 'Tags' in vpcd1:
                for tags in vpcd1['Tags']:
                    Name = tags['Value']
                    # print(Name)
                    vpcd1['VpcName'] = Name
            else:
                vpcd1['VpcName'] = 'No name'

            response1 = ec2_client.describe_subnets()
            sub = []
            for subd in response1['Subnets']:
                if VpcID == subd['VpcId']:
                    subnet = subd['SubnetId']
                    sub.append(subnet)
                    vpcd1['region'] = k
            vpcd1['subnet'] = sub
            response2 = ec2_client.describe_route_tables()
            for route in response2['RouteTables']:
                response2['RouteTabels'] = route['RouteTableId']
                if VpcID == route['VpcId']:
                    routetable = route['RouteTableId']
                    vpcd1['Routetables'] = routetable
            response3 = ec2_client.describe_internet_gateways()
            for igw in response3['InternetGateways']:
                # response3['InternetGateways'] = igw['InternetGatewayId']
                for att in igw['Attachments']:
                    if VpcID == att['VpcId']:
                        internetg = igw['InternetGatewayId']
                        vpcd1['InternetGateway'] = internetg
                        for ttg in igw['Tags']:
                            nname = ttg['Value']
                            vpcd1['igname'] = nname
                    else:
                        vpcd1['InternetGateway'] = ' - '
                        vpcd1['igname'] = " - "
            data.append(vpcd1)
    #         serial = vpc(data=vpcd1)
    #         if serial.is_valid():
    #             data.append(serial.data)
    #         else:
    #             print(serial.errors)
    # return data
    doc = docx.Document()

    sections = doc.sections
    for section in sections:
        section.top_margin = Mm(15)
        section.bottom_margin = Mm(15)
        section.left_margin = Mm(15)
        section.right_margin = Mm(15)
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(
        '                                                                                                                                                                           ')
    picture = run.add_picture('cloudjournee.png', width=Inches(1.25))

    para = doc.add_paragraph().add_run('AWS Inventory - Current Consolidated Report')

    para.font.size = Pt(24)
    font = para.font
    font.color.rgb = RGBColor(0, 0, 255)
    p = doc.add_paragraph('The following line are the list of assets audited')
    add_hyperlink(p, 'Link to my site', "https://www.cloudjournee.com/")
    para = doc.add_paragraph().add_run('VPC Summary')
    para.font.size = Pt(18)
    font = para.font
    font.color.rgb = RGBColor(255, 0, 0)
    menutable = doc.add_table(rows=1, cols=9)
    menutable.style = 'Colorful List'
    hdr_cells = menutable.rows[0].cells
    hdr_cells[0].text = 'VpcId'
    hdr_cells[1].text = 'VpcName'
    hdr_cells[2].text = 'CidrBlock'
    hdr_cells[3].text = 'IsDefault'
    hdr_cells[4].text = 'State'
    hdr_cells[5].text = 'subnet'
    hdr_cells[6].text = 'Routetables'
    hdr_cells[7].text = 'InternetGateway'
    hdr_cells[8].text = 'igname'

    for i in data:
        # print(i)
        row_cells = menutable.add_row().cells
        try:
            row_cells[0].text = i['VpcId']
            row_cells[1].text = i['VpcName']
            row_cells[2].text = i['CidrBlock']
            row_cells[3].text = str(i['IsDefault'])
            row_cells[4].text = i['State']
            row_cells[5].text = ",\n".join(i['subnet'])
            row_cells[6].text = i['Routetables']
            row_cells[7].text = i['InternetGateway']
            row_cells[8].text = i['igname']

        except:
            row_cells[8].text = ' - '

            for row in menutable.rows:
                for cell in row.cells:
                    paragraphs = cell.paragraphs
                    for paragraph in paragraphs:
                        for run in paragraph.runs:
                            font = run.font
                            font.size = Pt(10)
    doc.save('vpc.docx')


vpcdd()
