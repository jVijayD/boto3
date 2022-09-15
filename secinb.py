import boto3
import docx
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Mm, Inches, RGBColor, Pt


def sinb():
    ec2client = boto3.client('ec2')
    regions = ec2client.describe_regions(
        AllRegions=False
    )
    for i in regions['Regions']:
        k = (i['RegionName'])
        ec2client = boto3.client('ec2', region_name=k)
        response = ec2client.describe_security_group_rules()
        for sgn in response['SecurityGroupRules']:
            rlid = sgn['SecurityGroupRuleId']
            gid = sgn['GroupId']
            isg = sgn['IsEgress']
            ipo = sgn['IpProtocol']
            sgd = sgn['FromPort']

            if not isg:

                if ipo == '-1':

                    if sgd == -1:
                        print('Inbound rule allows all traffic', gid, rlid, k)

                        doc = docx.Document()

                        sections = doc.sections
                        for section in sections:
                            section.top_margin = Mm(10)
                            section.bottom_margin = Mm(10)
                            section.left_margin = Mm(10)
                            section.right_margin = Mm(10)
                        paragraph = doc.add_paragraph()
                        run = paragraph.add_run(
                            '                                                                                                                                                                           ')
                        run.add_picture('cloudjournee.png', width=Inches(1.25))

                        para = doc.add_paragraph().add_run('AWS Inventory - Current Consolidated Report')

                        para.font.size = Pt(24)
                        font = para.font
                        font.color.rgb = RGBColor(0, 0, 255)
                        menutable = doc.add_table(rows=2, cols=4)

                        menutable.style = 'Table Grid'

                        shading_elm_1 = parse_xml(r'<w:shd {} w:fill="1F5C8B"/>'.format(nsdecls('w')))
                        menutable.rows[0].cells[0]._tc.get_or_add_tcPr().append(shading_elm_1)
                        shading_elm_1 = parse_xml(r'<w:shd {} w:fill="1F5C8B"/>'.format(nsdecls('w')))
                        menutable.rows[0].cells[1]._tc.get_or_add_tcPr().append(shading_elm_1)
                        shading_elm_1 = parse_xml(r'<w:shd {} w:fill="1F5C8B"/>'.format(nsdecls('w')))
                        menutable.rows[0].cells[2]._tc.get_or_add_tcPr().append(shading_elm_1)
                        shading_elm_1 = parse_xml(r'<w:shd {} w:fill="1F5C8B"/>'.format(nsdecls('w')))
                        menutable.rows[0].cells[3]._tc.get_or_add_tcPr().append(shading_elm_1)
                        shading_elm_1 = parse_xml(r'<w:shd {} w:fill="#FF0000"/>'.format(nsdecls('w')))
                        menutable.rows[1].cells[2]._tc.get_or_add_tcPr().append(shading_elm_1)
                        hdr_cells = menutable.rows[0].cells
                        hdr_cells[0].text = 'AssetType'
                        hdr_cells[1].text = 'Risk/Potential Gap'
                        hdr_cells[2].text = 'Severity'
                        hdr_cells[3].text = 'Description'

                        row_cells = menutable.rows[1].cells
                        row_cells[0].text = ['SecurityGroup']
                        row_cells[1].text = [' Security Groups Inbound Rules Allows All Traffic']
                        row_cells[2].text = ['High']
                        row_cells[3].text = [
                            'It is recommended to not allow Inbound rules traffic to everyone. This should be immediately blocked.']

                        for row in menutable.rows:
                            for cell in row.cells:
                                paragraphs = cell.paragraphs
                                for paragraph in paragraphs:
                                    for run in paragraph.runs:
                                        font = run.font
                                        font.size = Pt(10)
                        widths = (Inches(0.5), Inches(2), Inches(0.5), Inches(4))
                        for row in menutable.rows:
                            for idx, width in enumerate(widths):
                                row.cells[idx].width = width
                        doc.save('secinb.docx')
                    else:
                        print('All traffic not allowed', gid, rlid, k)
                else:
                    print('Inbound TCP port or UDP port or ICMP')
            else:
                print('Outbound Traffic', gid, rlid, k)


sinb()
