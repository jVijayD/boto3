from datetime import date

import boto3
import docx
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Inches, Pt, RGBColor, Mm


def pas():
    paclient = boto3.client('iam')
    response = paclient.list_users()
    for udd in response['Users']:
        usr = udd['UserName']
        response = paclient.get_login_profile(UserName=usr)
        # rsp = response['LoginProfile']
        crtd = response['LoginProfile']['CreateDate'].date()
        currentdate = date.today()
        age = currentdate - crtd
        # print(age, usr)
        if age.days > 90:
            print(' IAM User password age is above 90 days', usr)

            doc = docx.Document()

            sections = doc.sections
            for section in sections:
                section.top_margin = Mm(10)
                section.bottom_margin = Mm(10)
                section.left_margin = Mm(10)
                section.right_margin = Mm(10)
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(
                '                                                                                                                                                                           ')
            run.add_picture('cloudjournee.png', width=Inches(1.25))

            para = doc.add_paragraph().add_run('AWS Inventory - Current Consolidated Report')

            para.font.size = Pt(24)
            font = para.font
            font.color.rgb = RGBColor(0, 0, 255)
            menutable = doc.add_table(rows=1, cols=4)

            menutable.style = 'Table Grid'

            shading_elm_1 = parse_xml(r'<w:shd {} w:fill="1F5C8B"/>'.format(nsdecls('w')))
            menutable.rows[0].cells[0]._tc.get_or_add_tcPr().append(shading_elm_1)
            shading_elm_1 = parse_xml(r'<w:shd {} w:fill="1F5C8B"/>'.format(nsdecls('w')))
            menutable.rows[0].cells[1]._tc.get_or_add_tcPr().append(shading_elm_1)
            shading_elm_1 = parse_xml(r'<w:shd {} w:fill="1F5C8B"/>'.format(nsdecls('w')))
            menutable.rows[0].cells[2]._tc.get_or_add_tcPr().append(shading_elm_1)
            shading_elm_1 = parse_xml(r'<w:shd {} w:fill="1F5C8B"/>'.format(nsdecls('w')))
            menutable.rows[0].cells[3]._tc.get_or_add_tcPr().append(shading_elm_1)
            shading_elm_1 = parse_xml(r'<w:shd {} w:fill="#FF0000"/>'.format(nsdecls('w')))

            hdr_cells = menutable.rows[0].cells
            hdr_cells[0].text = 'AssetType'
            hdr_cells[1].text = 'Risk/Potential Gap'
            hdr_cells[2].text = 'Severity'
            hdr_cells[3].text = 'Description'

            row_cells = menutable.add_row().cells
            # row_cells = menutable.rows[1].cells
            row_cells[0].text = ['IAM']
            row_cells[1].text = ['Checks if IAM User – Password Age/Policy']
            row_cells[2].text = ['High']
            row_cells[3].text = ['It is recommended that IAM User needs to adhere the AWS IAM password policies & AWS IAM password age period.']
            menutable.rows[1].cells[2]._tc.get_or_add_tcPr().append(shading_elm_1)
            for row in menutable.rows:
                for cell in row.cells:
                    paragraphs = cell.paragraphs
                    for paragraph in paragraphs:
                        for run in paragraph.runs:
                            font = run.font
                            font.size = Pt(10)
            widths = (Inches(0.5), Inches(2), Inches(0.5), Inches(4))
            for row in menutable.rows:
                for idx, width in enumerate(widths):
                    row.cells[idx].width = width
            doc.save('page.docx')
        else:
            print('is not above 90 days', usr)

    # response = paclient.get_account_password_policy()
    # vdd = response['PasswordPolicy']
    # page = response['MaxPasswordAge']


pas()
