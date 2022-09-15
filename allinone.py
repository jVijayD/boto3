import boto3
import docx
from docx.shared import Mm, Inches, Pt, RGBColor

data = []
ssn = []
llam = []
ecc = []
rss = []
eb = []



def vpcdd():
    ec2client = boto3.client('ec2')
    regions = ec2client.describe_regions(
        AllRegions=False
    )
    for i in regions['Regions']:
        k = (i['RegionName'])
        ec2_client = boto3.client('ec2', region_name=k)
        response = ec2_client.describe_vpcs()
        for vpcd1 in response['Vpcs']:
            VpcID = vpcd1['VpcId']
            if 'Tags' in vpcd1:
                for tags in vpcd1['Tags']:
                    Name = tags['Value']
                    # print(Name)
                    vpcd1['VpcName'] = Name
            else:
                vpcd1['VpcName'] = 'No name'

            response1 = ec2_client.describe_subnets()
            sub = []
            for subd in response1['Subnets']:
                if VpcID == subd['VpcId']:
                    subnet = subd['SubnetId']
                    sub.append(subnet)
                    vpcd1['region'] = k
            vpcd1['subnet'] = sub
            response2 = ec2_client.describe_route_tables()
            for route in response2['RouteTables']:
                response2['RouteTabels'] = route['RouteTableId']
                if VpcID == route['VpcId']:
                    routetable = route['RouteTableId']
                    vpcd1['Routetables'] = routetable
            response3 = ec2_client.describe_internet_gateways()
            for igw in response3['InternetGateways']:
                # response3['InternetGateways'] = igw['InternetGatewayId']
                for att in igw['Attachments']:
                    if VpcID == att['VpcId']:
                        internetg = igw['InternetGatewayId']
                        vpcd1['InternetGateway'] = internetg
                        for ttg in igw['Tags']:
                            nname = ttg['Value']
                            vpcd1['igname'] = nname
                    else:
                        vpcd1['InternetGateway'] = ' - '
                        vpcd1['igname'] = " - "
            data.append(vpcd1)
    snss()


def snss():
    ec2client = boto3.client('ec2')
    regions = ec2client.describe_regions(
        AllRegions=False
    )
    for i in regions['Regions']:
        k = (i['RegionName'])
        s_client = boto3.client('sns', region_name=k)
        response = s_client.list_subscriptions()
        for sss in response['Subscriptions']:
            sun = sss['SubscriptionArn']
            sss['region'] = k
            ssn.append(sss)

    ec2data()


def ec2data():
    ec2client = boto3.client('ec2')
    regions = ec2client.describe_regions(
        AllRegions=False
    )
    for i in regions['Regions']:
        k = (i['RegionName'])
        ec2_client = boto3.client('ec2', region_name=k)

        response1 = ec2_client.describe_images(Owners=[
            'self'
        ])
        for sna in response1['Images']:
            # print(k, sna)
            nname = sna['Name']
            img = sna['ImageId']

            # print(img)
            for bd in sna['BlockDeviceMappings']:

                try:
                    bd['SnapshotId'] = bd['Ebs']['SnapshotId']
                    bbb = bd['SnapshotId']
                    sna['SnapshotId'] = bbb

                    # print(nname, bbb, k)
                except:
                    print('  ')
            response = ec2_client.describe_instances(Filters=[
                {
                    'Name': 'image-id',
                    'Values': [
                        img,
                    ]
                },
            ], )
            sna['RegionName'] = k
            for ec2da in response['Reservations']:
                # print(ec2da)
                for instances in ec2da['Instances']:
                    insd = instances['InstanceId']
                    # print(insd)
                    sna['InstanceId'] = insd
            ecc.append(sna)

            lambdas()


def lambdas():
    ec2client = boto3.client('ec2')
    regions = ec2client.describe_regions(
        AllRegions=False
    )
    for i in regions['Regions']:
        k = (i['RegionName'])
        l_client = boto3.client('lambda', region_name=k)
        response = l_client.list_functions()
        for lll in response['Functions']:
            fun = lll['FunctionName']
            lll['region'] = k
            llam.append(lll)

        ebbs()


def ebbs():
    ec2client = boto3.client('ec2')
    regions = ec2client.describe_regions(
        AllRegions=False
    )
    for i in regions['Regions']:
        k = (i['RegionName'])
        ebclient = boto3.client('ec2', region_name=k)
        response = ebclient.describe_volumes()
        for eee in response['Volumes']:
            for ins in eee['Attachments']:
                bun = ins['InstanceId']
                eee['InstanceId'] = bun
                eee['region'] = k
                eb.append(eee)

        rdds()


def rdds():
    ec2client = boto3.client('ec2')
    regions = ec2client.describe_regions(
        AllRegions=False
    )
    for i in regions['Regions']:
        k = (i['RegionName'])
    rdclient = boto3.client('rds', region_name=k)
    response = rdclient.describe_db_instances()
    for rrr in response['DBInstances']:
        run = rrr['DBInstanceIdentifier']
        rrn = rrr['DBSubnetGroup']['VpcId']
        rrr['VpcId'] = rrn
        rrr['region'] = k
        rss.append(rrr)

    docs()


def docs():
    doc = docx.Document()

    sections = doc.sections
    for section in sections:
        section.top_margin = Mm(15)
        section.bottom_margin = Mm(15)
        section.left_margin = Mm(15)
        section.right_margin = Mm(15)
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(
        '                                                                                                                                                                           ')
    picture = run.add_picture('cloudjournee.png', width=Inches(1.25))

    para = doc.add_paragraph().add_run('AWS Inventory - Current Consolidated Report')

    para.font.size = Pt(24)
    font = para.font
    font.color.rgb = RGBColor(0, 0, 255)
    p = doc.add_paragraph('The following line are the list of assets audited')
    para = doc.add_paragraph().add_run('VPC Summary')
    para.font.size = Pt(18)
    font = para.font
    font.color.rgb = RGBColor(255, 0, 0)
    menutable = doc.add_table(rows=1, cols=9)
    menutable.style = 'Colorful List'
    hdr_cells = menutable.rows[0].cells
    hdr_cells[0].text = 'VpcId'
    hdr_cells[1].text = 'VpcName'
    hdr_cells[2].text = 'CidrBlock'
    hdr_cells[3].text = 'IsDefault'
    hdr_cells[4].text = 'State'
    hdr_cells[5].text = 'subnet'
    hdr_cells[6].text = 'Routetables'
    hdr_cells[7].text = 'InternetGateway'
    hdr_cells[8].text = 'igname'

    for i in data:
        # print(i)
        row_cells = menutable.add_row().cells
        try:
            row_cells[0].text = i['VpcId']
            row_cells[1].text = i['VpcName']
            row_cells[2].text = i['CidrBlock']
            row_cells[3].text = str(i['IsDefault'])
            row_cells[4].text = i['State']
            row_cells[5].text = i['subnet']
            row_cells[6].text = i['Routetables']
            row_cells[7].text = i['InternetGateway']
            row_cells[8].text = i['igname']
            for row in menutable.rows:
                for cell in row.cells:
                    paragraphs = cell.paragraphs
                    for paragraph in paragraphs:
                        for run in paragraph.runs:
                            font = run.font
                            font.size = Pt(8)
        except:
            row_cells[8].text = ' - '

            for row in menutable.rows:
                for cell in row.cells:
                    paragraphs = cell.paragraphs
                    for paragraph in paragraphs:
                        for run in paragraph.runs:
                            font = run.font
                            font.size = Pt(8)

    para = doc.add_paragraph().add_run('SNS Summary')
    para.font.size = Pt(18)
    font = para.font
    font.color.rgb = RGBColor(255, 0, 0)
    menutable = doc.add_table(rows=1, cols=5)
    menutable.style = 'Colorful List'
    hdr_cells = menutable.rows[0].cells
    hdr_cells[0].text = 'TopicArn'
    hdr_cells[1].text = 'region'
    hdr_cells[2].text = 'SubscriptionArn'
    hdr_cells[3].text = 'Protocol'
    hdr_cells[4].text = 'Endpoint'
    for j in ssn:
        row_cells = menutable.add_row().cells
        row_cells[0].text = j['TopicArn']
        row_cells[1].text = j['region']
        row_cells[2].text = j['SubscriptionArn']
        row_cells[3].text = j['Protocol']
        row_cells[4].text = j['Endpoint']

        for row in menutable.rows:
            for cell in row.cells:
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    for run in paragraph.runs:
                        font = run.font
                        font.size = Pt(8)

    para = doc.add_paragraph().add_run('AMI Summary')
    para.font.size = Pt(18)
    font = para.font
    font.color.rgb = RGBColor(255, 0, 0)
    menutable = doc.add_table(rows=1, cols=5)
    menutable.style = 'Colorful List'
    hdr_cells = menutable.rows[0].cells
    hdr_cells[0].text = 'ImageName'
    hdr_cells[1].text = 'ImageId'
    hdr_cells[2].text = 'SnapshotId'
    hdr_cells[3].text = 'InstanceId'
    hdr_cells[4].text = 'region'

    for n in data:
        # print(i)
        row_cells = menutable.add_row().cells
        try:
            row_cells[0].text = n['Name']
            row_cells[1].text = n['ImageId']
            row_cells[2].text = n['SnapshotId']
            row_cells[4].text = n['RegionName']
            row_cells[3].text = n['InstanceId']

        except:
            row_cells[3].text = ['  -   ']

    for row in menutable.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size = Pt(8)

    para = doc.add_paragraph().add_run('Lambda Summary')
    para.font.size = Pt(18)
    font = para.font
    font.color.rgb = RGBColor(255, 0, 0)
    menutable = doc.add_table(rows=1, cols=4)
    menutable.style = 'Colorful List'
    hdr_cells = menutable.rows[0].cells
    hdr_cells[0].text = 'FunctionName'
    hdr_cells[1].text = 'Runtime'
    hdr_cells[2].text = 'Role'
    hdr_cells[3].text = 'region'
    for m in llam:
        # print(i)
        row_cells = menutable.add_row().cells
        row_cells[0].text = m['FunctionName']
        row_cells[1].text = str(m['Runtime'])
        row_cells[2].text = m['Role']
        row_cells[3].text = m['region']

        for row in menutable.rows:
            for cell in row.cells:
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    for run in paragraph.runs:
                        font = run.font
                        font.size = Pt(8)

    para = doc.add_paragraph().add_run('EBS Summary')
    para.font.size = Pt(18)
    font = para.font
    font.color.rgb = RGBColor(255, 0, 0)
    menutable = doc.add_table(rows=1, cols=8)
    menutable.style = 'Colorful List'
    hdr_cells = menutable.rows[0].cells
    hdr_cells[0].text = 'VolumeId'
    hdr_cells[1].text = 'VolumeType'
    hdr_cells[2].text = 'Size'
    hdr_cells[3].text = 'CreateTime'
    hdr_cells[4].text = 'AvailabilityZone'
    hdr_cells[5].text = 'region'
    hdr_cells[6].text = 'State'
    hdr_cells[7].text = 'InstanceId'
    for b in eb:

        row_cells = menutable.add_row().cells
        row_cells[0].text = b['VolumeId']
        row_cells[1].text = b['VolumeType']
        row_cells[2].text = str(b['Size'])
        row_cells[3].text = str(b['CreateTime'])
        row_cells[4].text = b['AvailabilityZone']
        row_cells[5].text = b['region']
        row_cells[6].text = b['State']
        row_cells[7].text = b['InstanceId']

        for row in menutable.rows:
            for cell in row.cells:
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    for run in paragraph.runs:
                        font = run.font
                        font.size = Pt(8)

    para = doc.add_paragraph().add_run('RDS Summary')
    para.font.size = Pt(18)
    font = para.font
    font.color.rgb = RGBColor(255, 0, 0)
    menutable = doc.add_table(rows=1, cols=11)
    menutable.style = 'Colorful List'
    hdr_cells = menutable.rows[0].cells
    hdr_cells[0].text = 'DBInstanceIdentifier'
    hdr_cells[1].text = 'DBInstanceClass'
    hdr_cells[2].text = 'Engine'
    hdr_cells[3].text = 'StorageType'
    hdr_cells[4].text = 'InstanceCreateTime'
    hdr_cells[5].text = 'DBInstanceStatus'
    hdr_cells[6].text = 'VpcId'
    hdr_cells[7].text = 'MultiAZ'
    hdr_cells[8].text = 'BackupRetentionPeriod'
    hdr_cells[9].text = 'AvailabilityZone'
    hdr_cells[10].text = 'region'
    for p in rss:
        # print(i)
        row_cells = menutable.add_row().cells
        row_cells[0].text = p['DBInstanceIdentifier']
        row_cells[1].text = p['DBInstanceClass']
        row_cells[2].text = p['Engine']
        row_cells[3].text = p['StorageType']
        row_cells[4].text = str(p['InstanceCreateTime'])
        row_cells[5].text = p['DBInstanceStatus']
        row_cells[6].text = p['VpcId']
        row_cells[7].text = str(p['MultiAZ'])
        row_cells[8].text = str(p['BackupRetentionPeriod'])
        row_cells[9].text = p['AvailabilityZone']
        row_cells[10].text = p['region']

        for row in menutable.rows:
            for cell in row.cells:
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    for run in paragraph.runs:
                        font = run.font
                        font.size = Pt(8)

    doc.save('allin.docx')


vpcdd()
