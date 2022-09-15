import boto3
import datetime

import docx
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Mm, Inches, Pt, RGBColor


def nod():
    nclient = boto3.client('iam')
    response = nclient.list_users()
    for udd in response['Users']:
        usr = udd['UserName']
        try:
            psl = udd['PasswordLastUsed']
        except:
            print('User has no Usage of IAM', usr)

            doc = docx.Document()

            sections = doc.sections
            for section in sections:
                section.top_margin = Mm(10)
                section.bottom_margin = Mm(10)
                section.left_margin = Mm(10)
                section.right_margin = Mm(10)
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(
                '                                                                                                                                                                           ')
            run.add_picture('cloudjournee.png', width=Inches(1.25))

            para = doc.add_paragraph().add_run('AWS Inventory - Current Consolidated Report')

            para.font.size = Pt(24)
            font = para.font
            font.color.rgb = RGBColor(0, 0, 255)
            menutable = doc.add_table(rows=1, cols=4)

            menutable.style = 'Table Grid'

            shading_elm_1 = parse_xml(r'<w:shd {} w:fill="1F5C8B"/>'.format(nsdecls('w')))
            menutable.rows[0].cells[0]._tc.get_or_add_tcPr().append(shading_elm_1)
            shading_elm_1 = parse_xml(r'<w:shd {} w:fill="1F5C8B"/>'.format(nsdecls('w')))
            menutable.rows[0].cells[1]._tc.get_or_add_tcPr().append(shading_elm_1)
            shading_elm_1 = parse_xml(r'<w:shd {} w:fill="1F5C8B"/>'.format(nsdecls('w')))
            menutable.rows[0].cells[2]._tc.get_or_add_tcPr().append(shading_elm_1)
            shading_elm_1 = parse_xml(r'<w:shd {} w:fill="1F5C8B"/>'.format(nsdecls('w')))
            menutable.rows[0].cells[3]._tc.get_or_add_tcPr().append(shading_elm_1)
            shading_elm_1 = parse_xml(r'<w:shd {} w:fill="#FF0000"/>'.format(nsdecls('w')))

            hdr_cells = menutable.rows[0].cells
            hdr_cells[0].text = 'AssetType'
            hdr_cells[1].text = 'Risk/Potential Gap'
            hdr_cells[2].text = 'Severity'
            hdr_cells[3].text = 'Description'

            row_cells = menutable.add_row().cells
            # row_cells = menutable.rows[1].cells
            row_cells[0].text = ['IAM']
            row_cells[1].text = ['Checks for No Usage IAM User']
            row_cells[2].text = ['High']
            row_cells[3].text = ['Always recommended to delete unwanted IAM users.']
            menutable.rows[1].cells[2]._tc.get_or_add_tcPr().append(shading_elm_1)
            for row in menutable.rows:
                for cell in row.cells:
                    paragraphs = cell.paragraphs
                    for paragraph in paragraphs:
                        for run in paragraph.runs:
                            font = run.font
                            font.size = Pt(10)
            widths = (Inches(0.5), Inches(2), Inches(0.5), Inches(4))
            for row in menutable.rows:
                for idx, width in enumerate(widths):
                    row.cells[idx].width = width
            doc.save('nousage.docx')


nod()
